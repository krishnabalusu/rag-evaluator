"""
Document Loader — Step 1 of the RAG pipeline.

Supports: PDF, DOCX, TXT, HTML/HTM, Markdown
Returns a list of Document objects with text + metadata.
"""

from __future__ import annotations

import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional

logger = logging.getLogger(__name__)


@dataclass
class Document:
    """Single unit of ingested content."""
    text: str
    metadata: dict = field(default_factory=dict)

    def __repr__(self) -> str:
        preview = self.text[:80].replace("\n", " ")
        return f"Document(chars={len(self.text)}, source={self.metadata.get('source', '?')!r}, preview={preview!r})"


# ── Per-format loaders ─────────────────────────────────────────────────────

def _load_txt(path: Path) -> List[Document]:
    text = path.read_text(encoding="utf-8", errors="replace")
    return [Document(text=text, metadata={"source": str(path), "doc_type": "txt", "file_name": path.name})]


def _load_md(path: Path) -> List[Document]:
    text = path.read_text(encoding="utf-8", errors="replace")
    return [Document(text=text, metadata={"source": str(path), "doc_type": "markdown", "file_name": path.name})]


def _load_pdf(path: Path) -> List[Document]:
    try:
        import pypdf
    except ImportError:
        raise ImportError("Install pypdf: pip install pypdf")

    docs: List[Document] = []
    with open(path, "rb") as f:
        reader = pypdf.PdfReader(f)
        for page_num, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                docs.append(Document(
                    text=text,
                    metadata={
                        "source": str(path),
                        "doc_type": "pdf",
                        "file_name": path.name,
                        "page_number": page_num,
                        "total_pages": len(reader.pages),
                    }
                ))
    return docs


def _load_docx(path: Path) -> List[Document]:
    try:
        import docx
    except ImportError:
        raise ImportError("Install python-docx: pip install python-docx")

    document = docx.Document(str(path))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    text = "\n".join(paragraphs)
    return [Document(
        text=text,
        metadata={"source": str(path), "doc_type": "docx", "file_name": path.name}
    )]


def _load_html(path: Path) -> List[Document]:
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        raise ImportError("Install beautifulsoup4: pip install beautifulsoup4")

    html = path.read_text(encoding="utf-8", errors="replace")
    soup = BeautifulSoup(html, "html.parser")
    # Remove script and style tags
    for tag in soup(["script", "style", "nav", "footer", "header"]):
        tag.decompose()
    text = soup.get_text(separator="\n")
    return [Document(
        text=text,
        metadata={"source": str(path), "doc_type": "html", "file_name": path.name}
    )]


# ── Dispatch table ─────────────────────────────────────────────────────────

_LOADERS = {
    ".txt":  _load_txt,
    ".md":   _load_md,
    ".pdf":  _load_pdf,
    ".docx": _load_docx,
    ".html": _load_html,
    ".htm":  _load_html,
}


# ── Public API ─────────────────────────────────────────────────────────────

def load_document(path: str | Path) -> List[Document]:
    """Load a single file. Returns list of Documents (PDFs yield one per page)."""
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")
    ext = path.suffix.lower()
    loader = _LOADERS.get(ext)
    if loader is None:
        raise ValueError(f"Unsupported file type: {ext!r}. Supported: {list(_LOADERS)}")
    docs = loader(path)
    logger.info("Loaded %d document(s) from %s", len(docs), path.name)
    return docs


def load_directory(
    folder: str | Path,
    recursive: bool = True,
    extensions: Optional[set] = None,
) -> List[Document]:
    """
    Load all supported documents from a folder.

    Args:
        folder:     Directory path to scan.
        recursive:  If True, scan subdirectories too.
        extensions: Whitelist of extensions. Defaults to all supported.
    """
    folder = Path(folder)
    if not folder.is_dir():
        raise NotADirectoryError(f"Not a directory: {folder}")

    if extensions is None:
        extensions = set(_LOADERS.keys())

    pattern = "**/*" if recursive else "*"
    all_docs: List[Document] = []
    failed: List[str] = []

    for file_path in sorted(folder.glob(pattern)):
        if file_path.suffix.lower() not in extensions:
            continue
        try:
            docs = load_document(file_path)
            all_docs.extend(docs)
        except Exception as exc:
            logger.warning("Failed to load %s: %s", file_path.name, exc)
            failed.append(str(file_path))

    logger.info(
        "Directory load complete — %d docs from %s (%d failed)",
        len(all_docs), folder, len(failed)
    )
    if failed:
        logger.warning("Failed files: %s", failed)

    return all_docs
